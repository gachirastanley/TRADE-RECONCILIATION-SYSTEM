
import streamlit as st
import pdfplumber
import pandas as pd
import re
import io
from docx import Document


def extractor_ui():
    st.title("Account & Share Codes Extractor")

    uploaded_file = st.file_uploader("Upload Render Report", type=["pdf"])

    # ✅ CLEAN ACCOUNT CODE
    def clean_code(code):
        code = code.replace(" ", "")
        code = re.sub(r"R$", "", code)       # remove trailing R
        code = re.sub(r"\d+$", "", code)     # remove trailing numbers
        return code.strip()

    # ✅ GROUPING KEY
    def get_group_key(code):
        if len(code) > 1:
            return code[1:]
        return code

    # =========================
    # ✅ MAIN PROCESSING
    # =========================
    if uploaded_file is not None:

        st.info("Extracting...")

        selected_codes = {}
        share_codes = set()

        try:
            with pdfplumber.open(uploaded_file) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()

                    if text:

                        # =========================
                        # ✅ ACCOUNT CODES
                        # =========================
                        acc_matches = re.findall(
                            r"Unknown Account Code\s+([A-Z0-9\s]+)",
                            text
                        )

                        acc_matches = [clean_code(m) for m in acc_matches]

                        for code in acc_matches:
                            key = get_group_key(code)

                            if key not in selected_codes:
                                selected_codes[key] = code

                        # =========================
                        # ✅ SHARE CODES
                        # =========================
                        share_matches = re.findall(
                            r"Analysis Code\s+([A-Z0-9\.]+)",
                            text
                        )

                        share_codes.update(share_matches)

            # =========================
            # ✅ ACCOUNT RESULTS
            # =========================
            final_account_codes = list(selected_codes.values())

            df_accounts = pd.DataFrame({
                "Account Code": final_account_codes
            })

            st.subheader("✅ Account Codes")
            st.dataframe(df_accounts, use_container_width=True)

            # =========================
            # ✅ SHARE RESULTS
            # =========================
            df_shares = pd.DataFrame({
                "Share Code": list(share_codes)
            })

            st.subheader("✅ Share Codes")
            st.dataframe(df_shares, use_container_width=True)

            # =========================
            # ✅ EXCEL DOWNLOAD
            # =========================
            excel_output = io.BytesIO()

            with pd.ExcelWriter(excel_output, engine='openpyxl') as writer:
                df_accounts.to_excel(writer, sheet_name="Account Codes", index=False)
                df_shares.to_excel(writer, sheet_name="Share Codes", index=False)

            st.download_button(
                "📥 Download Excel",
                data=excel_output.getvalue(),
                file_name
=f"{base_name}.xlsx",
 

                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

            # =========================
            # ✅ WORD DOWNLOAD
            # =========================
            doc = Document()

            doc.add_heading('Account Codes', 0)
            for code in final_account_codes:
                doc.add_paragraph(code)

            doc.add_heading('Share Codes', 0)
            for code in share_codes:
                doc.add_paragraph(code)

            word_output = io.BytesIO()
            doc.save(word_output)

            st.download_button(
                "📄 Download Word",
                data=word_output.getvalue(),
                file_name
=f"{base_name}.docx",
 
xmlformats-officedocument.wordprocessingml.document"
            )



        except Exception as e:
            st.error(f"Error: {e}")
