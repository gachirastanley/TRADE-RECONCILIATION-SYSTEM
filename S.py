import streamlit as st
import pandas as pd
import os
import json
from datetime import datetime
from io import BytesIO
from table import extractor_ui

from openpyxl.styles import Font
from docx import Document
from docx.shared import Pt



from script import cdc_receipting_ui
from login import login_user, register_user


if "logged_in" not in st.session_state:
    st.session_state.logged_in = False

# =====================================================
# CONSTANTS
# =====================================================
HISTORY_FILE = "history.json"
HISTORY_FOLDER = "history_files"
os.makedirs(HISTORY_FOLDER, exist_ok=True)

# =====================================================
# GLOBAL STYLES
# =====================================================


def apply_styles():
    st.markdown("""
    <style>

    /* ✅ MAIN APP */
    .stApp {
        background-color: #ffffff;
        color: #0f172a;
        font-family: system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
    }

    /* ✅ TRANSPARENT BACKGROUND LOGO (WATERMARK) */
    .stApp::before {
        content: "";
        position: fixed;
        top: 50%;
        left: 50%;
        width: 320px;
        height: 320px;

        transform: translate(-50%, -50%);

        background-image: url("https://tse4.mm.bng.net/th/id/OIP.m9NFEGw194N3YNNJq4OO8gAAAA?rs=1&pid=ImgDetMain&o=7&rm=3");
        background-size: contain;
        background-repeat: no-repeat;

        opacity: 0.08;  /* ✅ much better watermark effect */
        z-index: -10;
    }

    /* ✅ CONTENT CONTAINER */
    .main > div {
        background: transparent;
        padding: 1.5rem 2rem;
    }

    /* ✅ HEADINGS */
    h1 { color: #0A2540; font-weight: 700; }
    h2, h3 { color: #102a43; font-weight: 600; }

    /* ✅ ANIMATED BUTTONS */
    div.stButton > button {
        background: linear-gradient(135deg, #1d4ed8, #1e40af);
        color: white;
        border-radius: 12px;
        padding: 0.55rem 1.4rem;
        font-weight: 600;
        border: none;
        cursor: pointer;

        box-shadow: 0 6px 16px rgba(29,78,216,0.35);
        transition: all 0.3s ease;
    }

    /* Hover */
    div.stButton > button:hover {
        transform: translateY(-3px) scale(1.03);
        box-shadow: 0 12px 30px rgba(29,78,216,0.5);
    }

    /* Click */
    div.stButton > button:active {
        transform: scale(0.96);
        box-shadow: 0 4px 10px rgba(29,78,216,0.3);
    }

    /* ✅ INPUT FIELDS */
    input, textarea, select {
        background-color: #ffffff !important;
        border: 1px solid #d1d5db !important;
        border-radius: 8px !important;
        color: #0f172a !important;
    }

    input:focus, textarea:focus {
        border-color: #2563eb !important;
        box-shadow: 0 0 0 2px rgba(37,99,235,0.2);
    }

    /* ✅ TABLE STYLING */
    .stDataFrame {
        background: #ffffff !important;
        border-radius: 12px;
        border: 1px solid #e5e7eb;
        box-shadow: 0 6px 16px rgba(15,23,42,0.08);
    }

    .stDataFrame thead tr th {
        background-color: #dbe4f1 !important;
        color: #102a43 !important;
        font-weight: 600;
    }

    .stDataFrame tbody tr td {
        background-color: #ffffff !important;
        color: #0f172a !important;
    }

    .stDataFrame tbody tr:hover td {
        background-color: #f1f5fb !important;
    }

    /* ✅ ALERTS */
    .stAlert {
        background: #ffffff !important;
        border: 1px solid #bfdbfe;
        border-radius: 10px;
        color: #1e3a8a;
        box-shadow: 0 4px 12px rgba(15,23,42,0.08);
    }

    /* ✅ DIVIDER */
    hr {
        border: none;
        height: 1px;
        background: #e5e7eb;
        margin: 1.5rem 0;
    }

    </style>
    """, unsafe_allow_html=True)


# ✅ CALL ONCE AT THE VERY TOP
apply_styles()


# =====================================================
# HISTORY HELPERS
# =====================================================
def load_history():
    if os.path.exists(HISTORY_FILE):
        with open(HISTORY_FILE, "r") as f:
            return json.load(f)
    return []


def save_history(data):
    with open(HISTORY_FILE, "w") as f:
        json.dump(data, f, indent=4)


# =====================================================
# UTILITY HELPERS
# =====================================================
def display_table_with_commas(df: pd.DataFrame, hide_index=False, hide_columns=False):
    styled = df.style.format(precision=2, thousands=",")
    styled = styled.apply(
        lambda row: ["font-weight: bold"] + [""] * (len(row) - 1),
        axis=1,
    )
    if hide_index:
        styled = styled.hide(axis="index")
    if hide_columns:
        styled = styled.hide(axis="columns")
    st.dataframe(styled, use_container_width=True)


def safe_find_col(df: pd.DataFrame, keywords):
    keywords = [k.lower() for k in keywords]
    for c in df.columns:
        if any(k in str(c).lower() for k in keywords):
            return c
    return None


def ensure_numeric(series: pd.Series):
    return pd.to_numeric(series, errors="coerce")


def sum_levies(df: pd.DataFrame, type_col: str, total_label: str) -> float:
    if df is None or type_col not in df.columns:
        return 0.0
    total_row = df[df[type_col] == total_label]
    if total_row.empty:
        return 0.0
    levy_keywords = ["broker", "stamp", "zse levy", "sec levy", "investor", "vat", "cgt"]
    total = 0.0
    for col in df.columns:
        if any(k in col.lower() for k in levy_keywords):
            val = pd.to_numeric(total_row.iloc[0][col], errors="coerce")
            if pd.notna(val):
                total += val
    return round(total, 2)


def extract_specific_levy(df: pd.DataFrame, type_col: str, total_label: str, keyword: str) -> float:
    if df is None or type_col not in df.columns:
        return 0.0
    total_row = df[df[type_col] == total_label]
    if total_row.empty:
        return 0.0
    for col in df.columns:
        if keyword in col.lower():
            val = pd.to_numeric(total_row.iloc[0][col], errors="coerce")
            return 0.0 if pd.isna(val) else float(val)
    return 0.0


# =====================================================
# SESSION STATE INIT
# =====================================================
st.session_state.setdefault("logged_in", False)
st.session_state.setdefault("username", "")
st.session_state.setdefault("sorted", False)
st.session_state.setdefault("zse_df", None)
st.session_state.setdefault("sh_df", None)
st.session_state.setdefault("show_final_summary", False)
st.session_state.setdefault("show_history", False)
st.session_state.setdefault("reconciled", False)
st.session_state.setdefault("show_extractor", False)
st.session_state.setdefault("show_audit", False)

# =====================================================
# LOGIN / REGISTER PAGE
# =====================================================

if "logged_in" not in st.session_state:
    st.session_state.logged_in = False

if not st.session_state.logged_in:

    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.image(
            "https://thumbs.dreamstime.com/b/psd-d-rendering-user-login-icon-psd-d-rendering-user-login-icon-isolated-transparent-background-357507015.jpg",
            width=300,
        )

    mode = st.toggle("Switch to Register")

    choice = "Register" if mode else "Login"

    if choice == "Login":
        st.subheader("Login")
        st.text("Please enter your details")
        username = st.text_input("Username", key="login_user")
        password = st.text_input("Password", type="password", key="login_pass")
        if st.button("Login", key="login_btn"):
            if login_user(username, password):
                st.session_state.logged_in = True
                st.session_state.username = username
                st.success(f"Welcome {username}")
                st.rerun()
            else:
                st.error("Invalid credentials")

    else:  # Register
        st.subheader("Register")
        new_user = st.text_input("New Username", key="reg_user")
        new_pass = st.text_input("New Password", type="password", key="reg_pass")
        if st.button("Register", key="reg_btn"):
            success, msg = register_user(new_user, new_pass)
            if success:
                st.success(msg)
            else:
                st.error(msg)

    st.stop()


# =====================================================
# MAIN APP (AUTHENTICATED)
# =====================================================
st.sidebar.success(f"Logged in as {st.session_state.username}")

st.sidebar.divider()

if st.sidebar.button("History"):
    st.session_state.show_history = True

if st.sidebar.button("Extractor"):
    st.session_state.show_extractor = True
    st.session_state.show_history = False

st.sidebar.divider()
if st.sidebar.button("Audit Trail"):
    st.session_state.show_audit = True
    st.session_state.show_history = False
    st.session_state.show_extractor = False

# push down
st.sidebar.markdown("<br>" * 22, unsafe_allow_html=True)

if st.sidebar.button("Logout"):
    st.session_state.logged_in = False
    st.session_state.username = ""
    st.rerun()


# =====================================================
# EXTRACTOR PAGE
# =====================================================
if st.session_state.show_extractor:

    extractor_ui()  # 👈 call your function from table.py

    if st.button("⬅ Back"):
        st.session_state.show_extractor = False
        st.rerun()

    st.stop()

# =====================================================
# HISTORY PAGE
# =====================================================
if st.session_state.show_history:
    allow_duplicates = st.checkbox("Allow duplicate reports", value=True)

    st.title("📜 History")
    history = load_history()

    if not history:
        st.info("No saved reports yet.")
    else:
        for item in list(reversed(history)):
            col1, col2, col3, col4, col5, col6 = st.columns(6)
            col1.write(item["date"])
            col2.write(item["user"])
            col3.write(item["type"])
            col4.write(f"{item['total']:.2f}")

            file_path = os.path.join(HISTORY_FOLDER, item["file"])

            if os.path.exists(file_path):
                with open(file_path, "rb") as f:
                    col5.download_button(
                        "Download",
                        data=f,
                        file_name=item["file"],
                        key=f"download_{item['file']}_{item['date']}",
                    )

            if col6.button("Delete", key=f"delete_{item['file']}_{item['date']}"):
                if os.path.exists(file_path):
                    os.remove(file_path)
                history.remove(item)
                save_history(history)
                st.success("Deleted successfully")
                st.rerun()

    if st.button("⬅ Back"):
        st.session_state.show_history = False
        st.rerun()

    st.stop()

# =====================================================
# AUDIT TRAIL PAGE
# =====================================================
if st.session_state.show_audit:

    st.title("🧾 Audit Trail")

    uploaded_file = st.file_uploader(
        "Upload Bank Statement File",
        type=["xlsx", "xls"]
    )

    if uploaded_file:

        import pandas as pd
        from openpyxl import load_workbook
        from io import BytesIO

        file_name = uploaded_file.name.lower()

        # =====================================================
        # ✅ 1. HANDLE FILE
        # =====================================================
        if file_name.endswith("xlsx"):
            wb = load_workbook(uploaded_file, data_only=True)

        elif file_name.endswith("xls"):

            st.info("🔄 Converting .xls to .xlsx...")

            try:
                df_temp = pd.read_excel(uploaded_file, engine="xlrd", header=None)
            except Exception:
                st.error("❌ Convert file to .xlsx and upload again.")
                st.stop()

            buffer = BytesIO()
            df_temp.to_excel(buffer, index=False, header=False)
            buffer.seek(0)

            wb = load_workbook(buffer, data_only=True)

        else:
            st.error("Unsupported file format")
            st.stop()

        ws = wb.active

        # =====================================================
        # ✅ 2. UNMERGE CELLS
        # =====================================================
        merged_ranges = list(ws.merged_cells.ranges)

        for merged in merged_ranges:
            min_row, min_col, max_row, max_col = merged.bounds
            value = ws.cell(row=min_row, column=min_col).value

            ws.unmerge_cells(str(merged))

            for r in range(min_row, max_row + 1):
                for c in range(min_col, max_col + 1):
                    ws.cell(row=r, column=c, value=value)

        # =====================================================
        # ✅ 3. CONVERT TO DATAFRAME
        # =====================================================
        df = pd.DataFrame(ws.values)

        # =====================================================
        # ✅ 4. SET ROW 13 AS HEADER
        # =====================================================
        header_row_index = 12

        if len(df) <= header_row_index:
            st.error("❌ Row 13 not found")
            st.stop()

        raw_headers = df.iloc[header_row_index]

        # ✅ MAKE UNIQUE HEADERS
        clean_headers = []
        seen = {}

        for h in raw_headers:
            name = str(h).strip() if pd.notna(h) else "EMPTY"

            if name in seen:
                seen[name] += 1
                name = f"{name}_{seen[name]}"
            else:
                seen[name] = 0

            clean_headers.append(name)

        df.columns = clean_headers
        df = df.iloc[header_row_index + 1:].reset_index(drop=True)

        # =====================================================
        # ✅ 5. FILTER DATA (FBC + INTERNET TRANSFERS)
        # =====================================================

        if len(df.columns) > 12:  # ensure enough columns

            col_k = df.columns[10]  # Column K
            col_b = df.columns[1]   # Column B
            col_m = df.columns[12]  # Column M

            cond_fbc = (
                df[col_k]
                .astype(str)
                .str.replace(r"\s+", " ", regex=True)  # normalize spaces
                .str.strip()
                .str.upper()
                .eq("FBC SECURITIES")
            )

            # ✅ Condition 2: INTERNET FUNDS TRANSFERS
            cond_internet = (
                    df[col_b]
                    .astype(str)
                    .str.upper()
                    .str.contains("INTERNET FUNDS TRSF BTWN ACCS", na=False)
                    &
                    df[col_m].notna()
            )

            # ✅ COMBINE CONDITIONS
            filtered_df = df[cond_fbc | cond_internet]

        else:
            st.error("❌ Required columns (B, K, M) not found")
            st.stop()
        # =====================================================
        # ✅ 6. FINAL CLEANING (KEY PART ✅)
        # =====================================================

        # ✅ Remove unwanted columns
        filtered_df = filtered_df.drop(
            columns=[
                c for c in filtered_df.columns
                if str(c).strip().lower() in ["deposit", "branch", "debit"]
            ],
            errors="ignore"
        )

        # ✅ Remove columns that are completely empty (NaN)
        filtered_df = filtered_df.dropna(axis=1, how='all')

        # ✅ Remove columns with only blank strings
        filtered_df = filtered_df.loc[:, ~(
            filtered_df.astype(str)
            .apply(lambda col: col.str.strip().eq("").all())
        )]

        # ✅ Remove columns with only "None"
        filtered_df = filtered_df.loc[:, ~(
            filtered_df.astype(str)
            .apply(lambda col: col.str.upper().eq("NONE").all())
        )]
        # =====================================================
        # ✅ 6.6 REMOVE LAST COLUMN
        # =====================================================
        if filtered_df.shape[1] > 0:
            filtered_df = filtered_df.iloc[:, :-1]
        # =====================================================
        # ✅ 6.5 FORMAT 'Trn Dt' AS DATE ONLY
        # =====================================================

        # find the column safely
        trn_col = next(
            (c for c in filtered_df.columns if "trn" in str(c).lower() and "dt" in str(c).lower()),
            None
        )

        if trn_col:
            filtered_df[trn_col] = pd.to_datetime(
                filtered_df[trn_col],
                errors="coerce"
            ).dt.date
        else:
            st.warning("⚠️ 'Trn Dt' column not found")

        # =====================================================
        # ✅ 7. DISPLAY
        # =====================================================
        st.success("✅ Bank Statement Data Processed Successfully")

        st.subheader("Bank Statement Data")
        st.dataframe(filtered_df, use_container_width=True)

        # ✅ Persist immediately so later sections (export, audit) never hit
        # a NameError if this block hasn't re-run on a given rerun.
        st.session_state.filtered_df = filtered_df

    # =====================================================
    # ✅ 8. UPLOAD & CLEAN SECOND FILE
    # =====================================================
    st.divider()

    second_file = st.file_uploader(
        "Upload FBC Bank Settlement file",
        type=["xlsx", "xls"],
        key="second_file"
    )

    if second_file:

        import pandas as pd

        try:
            # Load raw (no headers)
            second_df = pd.read_excel(second_file, header=None, engine="openpyxl")
        except Exception:
            try:
                second_df = pd.read_excel(second_file, header=None, engine="xlrd")
            except Exception:
                st.error("❌ Unable to read second file")
                st.stop()

        # =====================================================
        # ✅ 1. SET ROW 20 AS HEADER (index 19)
        # =====================================================
        header_row_index = 19

        if len(second_df) <= header_row_index:
            st.error("❌ Row 20 not found in file")
            st.stop()

        raw_headers = second_df.iloc[header_row_index]

        # ✅ Clean + make unique headers
        clean_headers = []
        seen = {}

        for h in raw_headers:
            name = str(h).strip() if pd.notna(h) else "EMPTY"

            if name in seen:
                seen[name] += 1
                name = f"{name}_{seen[name]}"
            else:
                seen[name] = 0

            clean_headers.append(name)

        second_df.columns = clean_headers

        # remove header rows above
        second_df = second_df.iloc[header_row_index + 1:].reset_index(drop=True)

        # =====================================================
        # ✅ 2. REMOVE EMPTY ROWS
        # =====================================================
        second_df = second_df.dropna(how="all")

        # remove rows where everything is blank string
        second_df = second_df[
            ~second_df.astype(str)
            .apply(lambda row: row.str.strip().eq("").all(), axis=1)
        ]

        # =====================================================
        # ✅ 3. REMOVE EMPTY COLUMNS
        # =====================================================
        second_df = second_df.loc[:, second_df.columns.notna()]
        second_df = second_df.loc[:, second_df.columns != "EMPTY"]

        # remove columns with all NaN
        second_df = second_df.dropna(axis=1, how="all")

        # remove columns with all blank strings
        second_df = second_df.loc[:, ~(
            second_df.astype(str)
            .apply(lambda col: col.str.strip().eq("").all())
        )]

        # remove columns with only "None"
        second_df = second_df.loc[:, ~(
            second_df.astype(str)
            .apply(lambda col: col.str.upper().eq("NONE").all())
        )]
        # =====================================================
        # ✅ 4. RENAME FIRST COLUMN TO "Date"
        # =====================================================
        if len(second_df.columns) > 0:
            cols = list(second_df.columns)
            cols[0] = "Date"
            second_df.columns = cols
        # =====================================================
        # ✅ 6. STANDARDIZE COLUMN NAMES
        # =====================================================

        expected_cols = ["Date", "Ref", "Type", "Details", "Debit", "Credit", "Balance"]

        # Trim or adjust based on available columns
        current_cols = list(second_df.columns)

        # If more columns exist → trim
        if len(current_cols) >= len(expected_cols):
            second_df = second_df.iloc[:, :len(expected_cols)]
            second_df.columns = expected_cols

        # If fewer columns exist → assign what is possible
        else:
            second_df.columns = expected_cols[:len(current_cols)]

        # =====================================================
        # ✅ 5. FINAL CLEAN OUTPUT
        # =====================================================
        st.success("✅ FBC Settlement data cleaned successfully")

        st.subheader("FBC Bank Settlement Data")
        st.dataframe(second_df, use_container_width=True)

        # store for later use
        st.session_state.second_df = second_df

    # =====================================================
    # ✅ 9. AUDIT TRAIL (FINAL CLEAN VERSION ✅)
    # =====================================================
    st.divider()

    if st.button("Audit Trail"):

        # Pull source frames from session_state so this works even if the
        # uploader blocks above didn't re-run this cycle.
        filtered_df = st.session_state.get("filtered_df")
        second_df = st.session_state.get("second_df")

        if filtered_df is None or second_df is None:
            st.error("❌ Please upload both the Bank Statement and the FBC Settlement file first.")
            st.stop()

        st.subheader("🔍 Full Reconciliation Results")

        # =====================================================
        # ✅ 1. GET REQUIRED COLUMNS
        # =====================================================
        bank_credit_col = next(
            (c for c in filtered_df.columns if "credit" in str(c).lower()),
            None
        )

        bank_date_col = next(
            (c for c in filtered_df.columns if "trn" in str(c).lower()),
            None
        )

        if bank_credit_col is None:
            st.error("❌ Bank credit column not found")
            st.stop()

        if not all(col in second_df.columns for col in ["Debit", "Credit", "Date"]):
            st.error("❌ Settlement columns missing")
            st.stop()

        # =====================================================
        # ✅ 2. CLEAN NUMERIC DATA
        # =====================================================
        filtered_df[bank_credit_col] = pd.to_numeric(
            filtered_df[bank_credit_col], errors="coerce"
        )

        second_df["Debit"] = pd.to_numeric(second_df["Debit"], errors="coerce")
        second_df["Credit"] = pd.to_numeric(second_df["Credit"], errors="coerce")

        # =====================================================
        # ✅ 3. TRACKING STRUCTURES
        # =====================================================
        matched_bank = set()
        matched_settlement = set()
        results = []

        # =====================================================
        # ✅ HELPER FUNCTION
        # =====================================================
        def extract_int(val):
            try:
                return int(float(val))
            except Exception:
                return None

        # =====================================================
        # ✅ MATCH LOGIC
        # =====================================================
        for i in range(len(second_df)):

            debit_val = second_df.iloc[i]["Debit"]

            if pd.isna(debit_val) or debit_val == 0:
                continue

            debit_int = extract_int(debit_val)

            direct_found = False

            # ✅ STEP 1: DIRECT MATCH
            for k, bank_row in filtered_df.iterrows():

                bank_val = bank_row[bank_credit_col]
                bank_int = extract_int(bank_val)

                if debit_int == bank_int:
                    matched_bank.add(k)
                    matched_settlement.add(i)

                    results.append({
                        "Date": bank_row[bank_date_col] if bank_date_col else "",
                        "Bank Credit": bank_val,
                        "Calculated": debit_val,
                        "Settlement Debit": debit_val,
                        "Settlement Credit": "",
                        "Match Type": "DIRECT",
                        "Result": "RECEIPTED ✅"
                    })

                    direct_found = True

            if direct_found:
                continue

            # ✅ STEP 2: CALCULATED MATCH
            for j in range(len(second_df)):

                credit_val = second_df.iloc[j]["Credit"]

                if pd.isna(credit_val) or credit_val == 0:
                    continue

                calculated = round(debit_val - credit_val, 2)
                calc_int = extract_int(calculated)

                for k, bank_row in filtered_df.iterrows():

                    if k in matched_bank:
                        continue

                    bank_val = bank_row[bank_credit_col]
                    bank_int = extract_int(bank_val)

                    if calc_int == bank_int:
                        matched_bank.add(k)
                        matched_settlement.add(i)
                        matched_settlement.add(j)

                        results.append({
                            "Date": bank_row[bank_date_col] if bank_date_col else "",
                            "Bank Credit": bank_val,
                            "Calculated": calculated,
                            "Settlement Debit": debit_val,
                            "Settlement Credit": credit_val,
                            "Match Type": "CALCULATED",
                            "Result": "RECEIPTED ✅"
                        })

        # =====================================================
        # ✅ 5. UNMATCHED BANK
        # =====================================================
        unmatched_bank = filtered_df.drop(index=list(matched_bank), errors="ignore")

        for _, row in unmatched_bank.iterrows():
            results.append({
                "Date": row[bank_date_col] if bank_date_col else "",
                "Bank Credit": row[bank_credit_col],
                "Settlement Debit": "",
                "Settlement Credit": "",
                "Match Type": "",
                "Result": "NOT RECEIPTED ❌"
            })

        # =====================================================
        # ✅ 6. UNMATCHED SETTLEMENT
        # =====================================================
        for idx, row in second_df.iterrows():
            if idx not in matched_settlement:

                if pd.notna(row["Debit"]) or pd.notna(row["Credit"]):
                    results.append({
                        "Date": row["Date"],
                        "Bank Credit": "",
                        "Settlement Debit": row["Debit"],
                        "Settlement Credit": row["Credit"],
                        "Match Type": "",
                        "Result": "NOT RECEIPTED ❌"
                    })

        # =====================================================
        # ✅ CLEAN DUPLICATES / CONFLICTS (FIXED)
        # =====================================================
        result_df = pd.DataFrame(results)

        RECEIPTED = "RECEIPTED ✅"
        NOT_RECEIPTED = "NOT RECEIPTED ❌"

        is_receipted = result_df["Result"] == RECEIPTED
        is_not_receipted = result_df["Result"] == NOT_RECEIPTED

        # ✅ Remove NOT-RECEIPTED bank rows whose value was actually matched elsewhere
        receipted_values = set(
            pd.to_numeric(result_df.loc[is_receipted, "Bank Credit"], errors="coerce").dropna()
        )

        result_df = result_df[
            ~(
                    is_not_receipted &
                    pd.to_numeric(result_df["Bank Credit"], errors="coerce").isin(receipted_values)
            )
        ]

        # refresh boolean masks after filtering
        is_receipted = result_df["Result"] == RECEIPTED
        is_not_receipted = result_df["Result"] == NOT_RECEIPTED

        # ✅ Remove NOT-RECEIPTED settlement rows whose debit was actually matched elsewhere
        receipted_debits = set(
            pd.to_numeric(result_df.loc[is_receipted, "Settlement Debit"], errors="coerce").dropna()
        )

        result_df = result_df[
            ~(
                    is_not_receipted &
                    pd.to_numeric(result_df["Settlement Debit"], errors="coerce").isin(receipted_debits)
            )
        ]

        # ✅ Final clean — keep ALL receipted rows AND all genuinely not-receipted rows
        # (no trimming/dropping below — every unmatched row in column 2 is preserved)
        result_df = result_df.drop_duplicates().reset_index(drop=True)

        # ✅ Persist for the export section below
        st.session_state.result_df = result_df
        st.session_state.filtered_df = filtered_df
        st.session_state.second_df = second_df

        # =====================================================
        # ✅ 7. FINAL OUTPUT
        # =====================================================
        st.success("✅ Full Reconciliation Complete")

        st.dataframe(result_df, use_container_width=True)

        # =====================================================
        # ✅ SUMMARY
        # =====================================================
        st.write("### Summary")
        st.write("✅ RECEIPTED:", (result_df["Result"] == RECEIPTED).sum())
        st.write("❌ NOT RECEIPTED:", (result_df["Result"] == NOT_RECEIPTED).sum())

    # =====================================================
    # ✅ EXPORT SECTION (only shown once an audit result exists)
    # =====================================================
    if (
        "result_df" in st.session_state
        and "filtered_df" in st.session_state
        and "second_df" in st.session_state
    ):

        def export_to_excel():
            result_df_local = st.session_state.result_df
            filtered_df_local = st.session_state.filtered_df
            second_df_local = st.session_state.second_df

            output = BytesIO()

            with pd.ExcelWriter(output, engine="openpyxl") as writer:

                filtered_df_local.to_excel(writer, sheet_name="Bank Data", index=False)
                second_df_local.to_excel(writer, sheet_name="Settlement Data", index=False)
                result_df_local.to_excel(writer, sheet_name="Audit Result", index=False)

                for ws in writer.book.worksheets:
                    for row in ws.iter_rows():
                        for cell in row:
                            cell.font = Font(name="Times New Roman", size=12)

                    for cell in ws[1]:
                        if cell.value:
                            cell.font = Font(name="Times New Roman", size=12, bold=True)

            output.seek(0)
            return output

        excel_file = export_to_excel()

        st.download_button(
            label="📥 Download Excel Report",
            data=excel_file,
            file_name="Audit_Trail_Report.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

        # =====================================================
        # ✅ EXPORT TO WORD
        # =====================================================
        def export_to_word():
            doc = Document()

            style = doc.styles['Normal']
            style.font.name = "Times New Roman"
            style.font.size = Pt(12)

            def add_table(df, title):
                doc.add_heading(title, level=2)

                table = doc.add_table(rows=1, cols=len(df.columns))

                # ✅ Header
                for i, col in enumerate(df.columns):
                    cell = table.rows[0].cells[i]
                    cell.text = str(col)
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
                            run.font.name = "Times New Roman"
                            run.font.size = Pt(12)

                # ✅ Data
                for _, row in df.iterrows():
                    cells = table.add_row().cells
                    for i, val in enumerate(row):
                        cells[i].text = str(val)
                        for p in cells[i].paragraphs:
                            for run in p.runs:
                                run.font.name = "Times New Roman"
                                run.font.size = Pt(12)

            # ✅ Add all tables (pulled from session_state, always defined here)
            add_table(st.session_state.filtered_df, "Bank Data")
            add_table(st.session_state.second_df, "Settlement Data")
            add_table(st.session_state.result_df, "Audit Results")

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer

        st.download_button(
            label="📄 Download Word Report",
            data=export_to_word(),
            file_name="Audit_Report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    # =====================================================
    # ✅ BACK BUTTON
    # =====================================================
    if st.button("⬅ Back"):
        st.session_state.show_audit = False
        st.rerun()

    st.stop()

# =====================================================
# MAIN RECEIPTING UI
# =====================================================
st.title("Trades Reconciliation System")
st.divider()


receipting_type = st.radio(
    "Select Receipting Type",
    ["ZSE Receipting", "CDC Receipting"],
    horizontal=True,
    key="receipting_radio",
)
st.divider()


# =====================================================
# CDC RECEIPTING
# =====================================================
if receipting_type == "CDC Receipting":
    cdc_receipting_ui()
    st.stop()


# =====================================================
# ZSE RECEIPTING
# =====================================================
import camelot  # noqa: E402 – imported here to avoid loading unless needed
# Top header with refresh
col1, col2 = st.columns([9, 1])

with col2:
    if st.button("🔄"):
        # Preserve login info
        keep_keys = ["logged_in", "username"]

        # Clear everything else
        for key in list(st.session_state.keys()):
            if key not in keep_keys:
                del st.session_state[key]

        st.rerun()
st.subheader("Start ZSE Reconciliation")


# ----- ZSE FILE UPLOAD -----
zse_file = st.file_uploader("Upload ZSE Excel File", type=["xlsx", "xls"])

if zse_file:
    raw_zse = pd.read_excel(zse_file, engine="openpyxl")
    raw_zse.columns = raw_zse.columns.str.lower().str.strip()

    buy_sell_col = safe_find_col(raw_zse, ["buy"])

    if buy_sell_col is None:
        st.error("Could not automatically detect a BUY/SELL column in the ZSE file.")
    else:
        if st.button("Sort ZSE"):
            df = raw_zse.copy()
            df[buy_sell_col] = df[buy_sell_col].astype(str).str.upper()

            numeric_cols = [
                c for c in df.select_dtypes(include="number").columns
                if str(c).lower().strip() != "due"
            ]

            def zse_totals(data, label):
                total = {
                    c: data[c].sum() if c in numeric_cols else ""
                    for c in data.columns
                }
                total[buy_sell_col] = label
                return pd.concat([data, pd.DataFrame([total])], ignore_index=True)

            st.session_state.zse_df = pd.concat(
                [
                    zse_totals(df[df[buy_sell_col] == "BUY"], "BUY TOTAL"),
                    zse_totals(df[df[buy_sell_col] == "SELL"], "SELL TOTAL"),
                ],
                ignore_index=True,
            )
            st.session_state.sorted = True

    if st.session_state.zse_df is not None:
        st.subheader("ZSE Data (with Totals)")
        st.dataframe(st.session_state.zse_df, use_container_width=True)


# ----- SHARESTOCK FILE UPLOAD -----
if st.session_state.sorted:

    st.divider()
    st.subheader("Sharestock Data")

    sh_file = st.file_uploader(
        "Upload Sharestock Excel File", type=["xlsx", "xls"], key="sh"
    )

    if sh_file:
        raw = pd.read_excel(sh_file, header=None, engine="openpyxl")

        header_idx = None
        for i, v in raw.iloc[:, 0].items():
            if isinstance(v, str) and v.strip().upper() == "CLIENT":
                header_idx = i
                break

        if header_idx is None:
            st.error("Could not find a 'CLIENT' header row in the Sharestock file.")
        else:
            headers = raw.loc[header_idx]
            clean_headers, seen = [], {}
            for i, h in enumerate(headers):
                name = str(h).strip() if pd.notna(h) and str(h).strip() else f"UNNAMED_{i}"
                if name in seen:
                    seen[name] += 1
                    name = f"{name}_{seen[name]}"
                else:
                    seen[name] = 0
                clean_headers.append(name)

            sh_df = raw.loc[header_idx + 1:].copy()
            sh_df.columns = clean_headers
            sh_df = sh_df.loc[:, ~sh_df.columns.str.startswith("UNNAMED")]
            sh_df = sh_df.dropna(how="all").reset_index(drop=True)

            qty_col = safe_find_col(sh_df, ["qty"])
            if qty_col is None:
                st.error("Could not find a 'Qty' column in the Sharestock data.")
            else:
                qty_idx = sh_df.columns.get_loc(qty_col)
                for col in sh_df.columns[qty_idx:]:
                    sh_df[col] = (
                        sh_df[col]
                        .astype(str)
                        .str.replace(",", "", regex=False)
                        .str.strip()
                    )
                    sh_df[col] = ensure_numeric(sh_df[col])

                st.session_state.sh_df = sh_df
                st.dataframe(sh_df, use_container_width=True)

        # ----- MATCH + RECONCILE -----
        if (
                st.button("Match Sharestock to ZSE")
                and st.session_state.zse_df is not None
                and st.session_state.sh_df is not None
        ):
            zse = st.session_state.zse_df.copy()
            sh = st.session_state.sh_df.copy()

            zse_type = safe_find_col(zse, ["buy", "sell"])
            zse_sym = safe_find_col(zse, ["symbol", "security", "counter"])
            sh_type = safe_find_col(sh, ["type"])
            sh_sym = safe_find_col(sh, ["symbol", "security"])

            if None in (zse_type, zse_sym, sh_type, sh_sym):
                st.error(
                    "Could not reliably detect transaction type / symbol columns in one of the files."
                )
            else:
                sh[sh_type] = sh[sh_type].replace({"BUY": "PURCHASE", "SELL": "SALE"})
                zse["_MATCH_"] = zse[zse_type].map({"BUY": "PURCHASE", "SELL": "SALE"})
                zse_clean = zse[~zse[zse_type].isin(["BUY TOTAL", "SELL TOTAL"])].copy()
                keys = set(zip(zse_clean[zse_sym], zse_clean["_MATCH_"]))

                def add_total_block(data, label):
                    total = {}
                    for col in data.columns:
                        if "price" in col.lower():
                            total[col] = ""
                        elif pd.api.types.is_numeric_dtype(data[col]):
                            total[col] = data[col].sum()
                        else:
                            total[col] = ""
                    total[sh_type] = label
                    return pd.concat([data, pd.DataFrame([total])], ignore_index=True)

                matched = sh[sh.apply(lambda r: (r[sh_sym], r[sh_type]) in keys, axis=1)]
                purchase_df = add_total_block(matched[matched[sh_type] == "PURCHASE"], "PURCHASE TOTAL")
                sale_df = add_total_block(matched[matched[sh_type] == "SALE"], "SALE TOTAL")

                column_mapping = {
                    "qty": ["quantity"],
                    "gross proceeds": ["gross consideration"],
                    "brokerage": ["broker", "commission"],
                    "basic": ["basic"],
                    "stamp": ["stamp"],
                    "zse levy": ["zse levy"],
                    "csd levy": ["depository levy"],
                    "cgt": ["capital gains tax"],
                    "sec levy": ["secz levy"],
                    "investor levy": ["inv protection vx"],
                    "vat": ["vat"],
                }

                def append_zse_row(df, zse_row, label):
                    row = {}
                    for sh_col in df.columns:
                        row[sh_col] = ""
                        if sh_col == sh_type:
                            row[sh_col] = label
                            continue
                        for sh_key, zse_keys in column_mapping.items():
                            if sh_key in sh_col.lower():
                                for zc in zse_row.index:
                                    if any(k in zc.lower() for k in zse_keys):
                                        row[sh_col] = zse_row[zc]
                    return pd.concat([df, pd.DataFrame([row])], ignore_index=True)

                zse_buy_total = zse[zse[zse_type] == "BUY TOTAL"]
                zse_sell_total = zse[zse[zse_type] == "SELL TOTAL"]

                if not zse_buy_total.empty:
                    purchase_df = append_zse_row(purchase_df, zse_buy_total.iloc[0], "ZSE BUY TOTAL")
                if not zse_sell_total.empty:
                    sale_df = append_zse_row(sale_df, zse_sell_total.iloc[0], "ZSE SELL TOTAL")

                def append_variance(df, total_label, zse_label, variance_label):
                    total_rows = df[df[sh_type] == total_label]
                    zse_rows = df[df[sh_type] == zse_label]
                    if total_rows.empty or zse_rows.empty:
                        return df
                    t = total_rows.iloc[0]
                    z = zse_rows.iloc[0]
                    row = {}
                    for col in df.columns:
                        if col == sh_type:
                            row[col] = variance_label
                        elif pd.api.types.is_numeric_dtype(df[col]):
                            tv = pd.to_numeric(t[col], errors="coerce")
                            zv = pd.to_numeric(z[col], errors="coerce")
                            row[col] = (0 if pd.isna(tv) else tv) - (0 if pd.isna(zv) else zv)
                        else:
                            row[col] = ""
                    return pd.concat([df, pd.DataFrame([row])], ignore_index=True)

                purchase_df = append_variance(
                    purchase_df, "PURCHASE TOTAL", "ZSE BUY TOTAL", "VARIANCE (PURCHASE - ZSE BUY)"
                )
                sale_df = append_variance(
                    sale_df, "SALE TOTAL", "ZSE SELL TOTAL", "VARIANCE (SALE - ZSE SELL)"
                )

                st.session_state.purchase_df = purchase_df
                st.session_state.sale_df = sale_df
                st.session_state.reconciled = True
                st.rerun()

            # ----- SHOW RECONCILED TABLES (persisted via session state) -----
        if st.session_state.get("reconciled") and "purchase_df" in st.session_state and "sale_df" in st.session_state:
            st.success("Full reconciliation and settlement summary complete")
            st.subheader("Sharestock – Purchases (Reconciled)")
            st.dataframe(st.session_state.purchase_df, use_container_width=True)
            st.subheader("Sharestock – Sales (Reconciled)")
            st.dataframe(st.session_state.sale_df, use_container_width=True)

            if st.button("Final Settlement Summary"):
                st.session_state.show_final_summary = True
                st.rerun()

            # =====================================================
            # FINAL SETTLEMENT SUMMARY
            # =====================================================
        if st.session_state.show_final_summary:

            st.subheader("Final Settlement Summary")

            bank_statement_amount = st.number_input(
                "Bank Statement Amount",
                min_value=0.0,
                value=0.0,
                step=0.01,
                format="%.2f",
            )

            # --- Compute values ---
            total_purchases_value = 0.0
            total_sales_value = 0.0
            zse_levy_remitted = 0.0
            ipl_levy_remitted = 0.0
            sec_levy_total = 0.0

            if "purchase_df" in st.session_state:
                p = st.session_state.purchase_df
                p_type = safe_find_col(p, ["type"])
                if p_type:
                    total_purchases_value = sum_levies(p, p_type, "PURCHASE TOTAL")
                    zse_levy_remitted += extract_specific_levy(p, p_type, "PURCHASE TOTAL", "zse levy")
                    ipl_levy_remitted += extract_specific_levy(p, p_type, "PURCHASE TOTAL", "investor")
                    sec_levy_total += extract_specific_levy(p, p_type, "PURCHASE TOTAL", "sec levy")

            if "sale_df" in st.session_state:
                s = st.session_state.sale_df
                s_type = safe_find_col(s, ["type"])
                if s_type:
                    total_sales_value = sum_levies(s, s_type, "SALE TOTAL")
                    zse_levy_remitted += extract_specific_levy(s, s_type, "SALE TOTAL", "zse levy")
                    ipl_levy_remitted += extract_specific_levy(s, s_type, "SALE TOTAL", "investor")
                    sec_levy_total += extract_specific_levy(s, s_type, "SALE TOTAL", "sec levy")

            total_purchases_and_sales = round(total_purchases_value + total_sales_value, 2)
            total_amount_to_be_received = round(
                total_purchases_and_sales - (zse_levy_remitted + ipl_levy_remitted + sec_levy_total), 2
            )
            balance_from_bank = round(bank_statement_amount - total_amount_to_be_received, 2)

            # Post settlement total from "due" column
            post_settlement_total = 0.0
            if st.session_state.zse_df is not None:
                due_col = next(
                    (c for c in st.session_state.zse_df.columns if str(c).lower().strip() == "due"),
                    None,
                )
                if due_col:
                    post_settlement_total = pd.to_numeric(
                        st.session_state.zse_df[due_col], errors="coerce"
                    ).sum()
            post_settlement_total = round(post_settlement_total, 2)
            difference_bank_vs_post_settlement = round(bank_statement_amount - post_settlement_total, 2)

            # Capital gains variance
            capital_gains_variance = 0.0
            if "sale_df" in st.session_state:
                s = st.session_state.sale_df
                s_type = safe_find_col(s, ["type"])
                if s_type:
                    vr = s[s[s_type] == "VARIANCE (SALE - ZSE SELL)"]
                    if not vr.empty:
                        cgt_cols = [c for c in vr.columns if "cgt" in c.lower()]
                        if cgt_cols:
                            val = pd.to_numeric(vr.iloc[0][cgt_cols[0]], errors="coerce")
                            capital_gains_variance = 0.0 if pd.isna(val) else round(float(val), 4)

            verify_value = round(
                bank_statement_amount + capital_gains_variance - post_settlement_total, 2
            )

            # --- Summary table ---
            summary_df = pd.DataFrame(
                [
                    ["Total Purchases", total_purchases_value, "", "", "", ""],
                    ["Total Sales", total_sales_value, "", "", "", ""],
                    ["Total Purchases and Sales", total_purchases_and_sales, "", "", "", ""],
                    ["", "", "", "", "", ""],
                    ["ZSE Levy Remitted Directly", zse_levy_remitted, "", "", "", ""],
                    ["IPL Levy Remitted Directly", ipl_levy_remitted, "", "", "", ""],
                    ["SEC Levy", sec_levy_total, "", "", "", ""],
                    [
                        "", "", "",
                        "Bank Statement Amount",
                        "Balance from Amount to be Received and Bank Amount",
                        "Capital Gains",
                        "Post Settlement Total from Report",
                        "Difference Between Bank and Post Settlement",
                        "Verify",
                    ],
                    [
                        "Total Amount to be Received",
                        total_amount_to_be_received,
                        "",
                        bank_statement_amount,
                        balance_from_bank,
                        capital_gains_variance,
                        post_settlement_total,
                        difference_bank_vs_post_settlement,
                        verify_value,
                    ],
                    ["", "", "", "", "", ""],
                    ["", "", "", "", "", ""],

                    ["Prepared by:....................................................................................................",
                     "", "", "", "", ""],
                    ["", "", "", "", "", ""],
                    ["Reviewed by:....................................................................................................",
                     "", "", "", "", ""],
                ]
            )
            display_table_with_commas(summary_df, hide_index=True, hide_columns=True)

            # --- Excel export ---
            from openpyxl.styles import Font, Alignment, Border, Side
            from openpyxl.utils import get_column_letter

            output = BytesIO()

            with pd.ExcelWriter(output, engine="openpyxl") as writer:

                st.session_state.sh_df.to_excel(writer, sheet_name="Sharestock Raw", index=False)
                st.session_state.purchase_df.to_excel(writer, sheet_name="Purchases", index=False)

                SALES_SHEET = "Sales & Summary"
                st.session_state.sale_df.to_excel(writer, sheet_name=SALES_SHEET, index=False, startrow=0)
                summary_df.to_excel(
                    writer,
                    sheet_name=SALES_SHEET,
                    index=False,
                    startrow=len(st.session_state.sale_df) + 3,
                    header=False,
                )

                bold = Font(bold=True)
                thick = Side(style="medium")
                thin = Side(style="thin")
                double_bottom = Side(style="double")

                # ---------- Purchases sheet ----------
                ws_purchases = writer.sheets["Purchases"]
                for cell in ws_purchases[1]:
                    cell.font = bold
                last_row = ws_purchases.max_row
                for r in range(max(1, last_row - 2), last_row + 1):
                    for c in range(1, ws_purchases.max_column + 1):
                        cell = ws_purchases.cell(row=r, column=c)
                        if "sharestock" not in str(cell.value or "").lower():
                            cell.font = bold

                # ---------- Sales & Summary sheet ----------
                ws_sales = writer.sheets[SALES_SHEET]

                # Bold header row
                for cell in ws_sales[1]:
                    cell.font = bold

                # Bold first empty row and below in column A
                first_empty_row = None
                for r in range(2, ws_sales.max_row + 1):
                    val = ws_sales.cell(row=r, column=1).value
                    if val is None or str(val).strip() == "":
                        first_empty_row = r
                        break
                if first_empty_row:
                    for r in range(first_empty_row, ws_sales.max_row + 1):
                        cell = ws_sales.cell(row=r, column=1)
                        if cell.value is None:
                            cell.value = ""
                        cell.font = bold

                # Bold last 3 rows that have values in col 17
                rows_with_values = []
                for r in range(ws_sales.max_row, 1, -1):
                    val = ws_sales.cell(row=r, column=17).value
                    if val not in (None, "", " "):
                        rows_with_values.append(r)
                    if len(rows_with_values) == 3:
                        break
                for r in rows_with_values:
                    for c in range(1, ws_sales.max_column + 1):
                        ws_sales.cell(row=r, column=c).font = bold

                # Thick border below "Total Sales" (cols A & B)
                for r in range(1, ws_sales.max_row + 1):
                    if str(ws_sales.cell(row=r, column=1).value or "").strip().lower() == "total sales":
                        ws_sales.cell(row=r, column=1).border = Border(bottom=thick)
                        ws_sales.cell(row=r, column=2).border = Border(bottom=thick)
                        break

                # Top + double-bottom border for "Total Amount to be Received"
                for r in range(1, ws_sales.max_row + 1):
                    if str(ws_sales.cell(row=r, column=1).value or "").strip().lower() == "total amount to be received":
                        ws_sales.cell(row=r, column=1).border = Border(top=thick, bottom=double_bottom)
                        ws_sales.cell(row=r, column=2).border = Border(top=thick, bottom=double_bottom)
                        break

                # Thick box around bank/verify header+value rows
                bank_headers = {
                    "bank statement amount",
                    "balance from amount to be received and bank amount",
                    "capital gains",
                    "post settlement total from report",
                    "difference between bank and post settlement",
                    "verify",
                }
                header_row = None
                header_cols = {}
                for r in range(1, ws_sales.max_row + 1):
                    for c in range(1, ws_sales.max_column + 1):
                        val = str(ws_sales.cell(row=r, column=c).value or "").strip().lower()
                        if val in bank_headers:
                            header_row = r
                            header_cols[val] = c
                    if header_row:
                        break

                if header_row and header_cols:
                    start_col = min(header_cols.values())
                    end_col = max(header_cols.values())
                    value_row = header_row + 1

                    for r in (header_row, value_row):
                        for c in range(start_col, end_col + 1):
                            ws_sales.cell(row=r, column=c).border = Border(
                                top=thick if r == header_row else None,
                                bottom=thick if r == value_row else None,
                                left=thick if c == start_col else thick,
                                right=thick if c == end_col else thick,
                            )

                    # Thin border on value row
                    for c in range(start_col, end_col + 1):
                        ws_sales.cell(row=value_row, column=c).border = Border(
                            top=thin, bottom=thin,
                            left=thin if c == start_col else thin,
                            right=thin if c == end_col else thin,
                        )

                # Wrap & centre row below "SEC Levy"
                for r in range(1, ws_sales.max_row + 1):
                    if str(ws_sales.cell(row=r, column=1).value or "").strip().lower() == "sec levy":
                        target_row = r + 1
                        if target_row <= ws_sales.max_row:
                            for c in range(1, ws_sales.max_column + 1):
                                ws_sales.cell(row=target_row, column=c).alignment = Alignment(
                                    wrap_text=True, horizontal="general", vertical="bottom"
                                )
                        break

                # Times New Roman 12 across all sheets
                for ws in writer.book.worksheets:
                    for row in ws.iter_rows():
                        for cell in row:
                            existing = cell.font or Font()
                            cell.font = Font(
                                name="Times New Roman",
                                size=12,
                                bold=existing.bold,
                                italic=existing.italic,
                                underline=existing.underline,
                                strike=existing.strike,
                                color=existing.color,
                            )

                # Auto column widths (all sheets)
                MAX_WIDTH, MIN_WIDTH, PADDING = 28, 8, 2
                # ✅ APPLY THOUSANDS FORMAT + ALIGNMENT
                num_fmt = '#,##0.00;-#,##0.00;-'

                for ws in writer.book.worksheets:
                    for row in ws.iter_rows():
                        for cell in row:

                            # ✅ Convert strings that look like numbers
                            if isinstance(cell.value, str):
                                try:
                                    cell.value = float(cell.value.replace(",", ""))
                                except Exception:
                                    pass

                            # ✅ Apply thousands separator
                            if isinstance(cell.value, (int, float)):
                                cell.number_format = num_fmt
                                cell.alignment = Alignment(
                                    horizontal="right",
                                    vertical="center"
                                )

                for ws in writer.book.worksheets:
                    for col_idx in range(1, ws.max_column + 1):
                        col_letter = get_column_letter(col_idx)
                        max_len = max(
                            (len(str(cell.value)) for cell in ws[col_letter] if cell.value is not None),
                            default=0,
                        )
                        ws.column_dimensions[col_letter].width = min(MAX_WIDTH, max(MIN_WIDTH, max_len + PADDING))

                writer.book.active = 0

            output.seek(0)
            file_name = f"ZSE_RECEIPTING_{post_settlement_total:.2f}.xlsx"
            file_path = os.path.join(HISTORY_FOLDER, file_name)

            if st.download_button(
                    label="📥 Save & Download Report",
                    data=output,
                    file_name=file_name,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ):
                with open(file_path, "wb") as f:
                    f.write(output.getvalue())

                history = load_history()
                new_entry = {
                    "file": file_name,
                    "user": st.session_state.username,
                    "date": datetime.now().strftime("%Y-%m-%d %H:%M"),
                    "type": receipting_type,
                    "total": float(post_settlement_total),
                }
                history.append(new_entry)

                save_history(history)
                st.success("CDC report saved to history.")
